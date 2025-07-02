from io import BytesIO

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from .. import crud, deps, schemas

router = APIRouter(prefix='/coverage', tags=['coverage'])


@router.get('/', response_model=list[schemas.CoverageRead])
def list_coverage(db: Session = Depends(deps.get_db)):
    return crud.get_all_coverage(db)


@router.post('/', response_model=schemas.CoverageCreate)
def create_coverage(
    data: schemas.CoverageCreate,
    db: Session = Depends(deps.get_db)
):
    return crud.create_coverage(db, data)


@router.get('/{coverage_id}', response_model=schemas.CoverageRead)
def get_coverage(coverage_id: int, db: Session = Depends(deps.get_db)):
    obj = crud.get_coverage(db, coverage_id)
    if not obj:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail='Coverage not found'
        )
    return obj


@router.get('/{coverage_id}/draft')
def draft_coverage_docx(coverage_id: int, db: Session = Depends(deps.get_db)):
    coverage = db.query(schemas.CoverageRead).filter(
        schemas.CoverageRead.id == coverage_id).first()
    if not coverage:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail='Coverage not found'
        )

    # === Create the DOCX ===
    doc = Document()
    doc.add_heading('Coverage Draft', 0)

    doc.add_paragraph(f'Coverage ID: {coverage.id}')
    doc.add_paragraph(f'Basis of Valuation: {coverage.basis_of_valuation}')
    doc.add_paragraph(f'Debit Note: {coverage.debit_note}')
    doc.add_paragraph(f'Ordinary Risks Rate: {coverage.ordinary_risks_rate}')
    doc.add_paragraph(f'War Risks Rate: {coverage.war_risks_rate}')
    doc.add_paragraph(f'Date: {coverage.date}')

    # === Save to memory buffer ===
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename=coverage_{coverage_id}_draft.docx'
        }
    )
